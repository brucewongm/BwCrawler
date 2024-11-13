import os
import time
from urllib.parse import urljoin
from pprint import pprint

import requests
from bs4 import BeautifulSoup
from translate import Translator
import re
from docx import Document

OPEN = 1
CLOSED = 0


def moment():
    return time.strftime("%Y-%m-%d_%H-%M-%S", time.localtime())


def english_to_chinese(text):
    translator = Translator(from_lang="en", to_lang="zh")
    reg = r'\s{2,}'
    pat = re.compile(reg, re.I | re.M)
    result = ''
    split_parts = re.split(pat, text)
    print(split_parts)
    for line in split_parts:
        line = line.strip()
        print('line:', line)
        print('mark start')
        if not line:
            print('mark empty')
            continue
            pass
        if len(line) >= 500:
            sentence_list = []
            second_units = line.split(',')
            for sentence in second_units:
                second_translation = translator.translate(sentence)
                sentence_list.append(second_translation)
                pass
            result += ','.join(sentence_list)
            pass
        else:
            translation = translator.translate(line)
            result += translation + '\n'
            pass
    return result


def txt2word(abs_full_txt_file_name, abs_full_word_file_name=None):
    # 读取TXT文件内容
    if not abs_full_word_file_name:
        abs_full_word_file_name = abs_full_txt_file_name.rsplit('.', maxsplit=1)[0] + '.docx'
    with open(abs_full_txt_file_name, mode='r', encoding='utf-8') as file:
        text = file.read()
        pass
    # 创建一个Document对象
    doc = Document()
    # 添加文本内容到Word文档
    doc.add_paragraph(text)
    # 保存Word文档
    doc.save(abs_full_word_file_name)
    pass


def pause(seconds: int):
    for i in range(seconds, 0, -1):
        time.sleep(1)
        print('\rpause counting down {}s'.format(i), end='')
        pass
    print('\rpause counting down 0s')
    pass


pass


class CrawlWomenSHealth:
    def __init__(self):
        self.mark_counter = 20241101
        pass

    def generate_mark(self):
        this_mark = 'mmmaaarrrkkk' + str(self.mark_counter)
        self.mark_counter += 1
        return this_mark

    def run(self):
        # 主页的URL
        # target_url = 'https://www.womenshealthmag.com/relationships/a46109633/what-is-a-swinger/'
        main_page_url = 'https://www.womenshealthmag.com/'
        refer_url = 'https://www.womenshealthmag.com/relationships/'
        relative_url = '/a45616686/how-to-tell-if-someone-likes-you/'
        full_url = urljoin(refer_url, relative_url)
        #
        result_txt_file_name = "women_s_health_{}english.txt".format(moment())
        current_folder = os.getcwd()
        parent_folder = os.path.dirname(current_folder)
        result_directory = os.path.join(parent_folder, 'results')
        result_abs_txt_file_name = os.path.join(result_directory, result_txt_file_name)
        if not os.path.exists(result_directory):
            os.makedirs(result_directory)
            pass
        pass

        # 设置请求头，模拟正常用户的浏览器请求
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/96.0.4664.45 Safari/537.36',
            'Accept-Language': 'en-US,en;q=0.5',
            'Accept-Encoding': 'gzip, deflate, br',
            # 'Connection': 'keep-alive',
            'Connection': 'close',
            'Upgrade-Insecure-Requests': '1',
            'DNT': '1',  # Do Not Track 请求头
            'Referer': refer_url,  # 引用页，有助于避免被识别为爬虫（有时）
        }
        # print('make-first-get,' * 10)
        pause(3)
        print('start requests getting......')
        # 获取cookies
        response_main_page = requests.get(main_page_url)
        # 获取Cookies
        cookies = response_main_page.cookies
        # 打印Cookies
        print('cookies:', cookies)
        # 或者将Cookies保存到一个cookiejar中，然后在后续请求中使用
        cookiejar = requests.cookies.RequestsCookieJar()
        for cookie in cookies:
            cookiejar.set(cookie.name, cookie.value)
            pass
        pass
        # print('make-second-get,' * 10)
        pause(10)
        #
        response = requests.get(full_url, headers=headers, cookies=cookiejar)
        print("response.status_code:")
        print(response.status_code)
        soup = BeautifulSoup(response.text, 'html.parser')
        #
        with open(result_abs_txt_file_name, mode='w', encoding='utf-8') as file:
            for element in soup.find_all(['h1', 'h2', 'h3', 'p', 'ul', 'li']):
                if element.name in ['h1', 'h2', 'h3', 'p']:
                    file.write(element.text.strip())
                    file.write('\n')
                    file.write(self.generate_mark())
                    file.write('\n')
                    file.flush()
                    pass
                elif element.name == 'ul':
                    list_items = element.find_all('li')
                    for li_element in list_items:
                        file.write('!!!:' + li_element.text.strip())
                        file.write('\n')
                        file.write(self.generate_mark())
                        file.write('\n')
                        file.flush()
                        pass
                    pass
                pass
            pass
        #
        txt2word(result_abs_txt_file_name, result_abs_txt_file_name.split('.')[0] + '.docx')
        pause(30)
        print('crawling finished!')
        pass

    pass


pass


def task1():
    ins = CrawlWomenSHealth()
    ins.run()
    pass


def compose_result_folder_doc_files():
    current_folder = os.getcwd()
    parent_folder = os.path.dirname(current_folder)
    target_folder = os.path.join(parent_folder, 'results')
    english_docx = ''
    chinese_docx = ''
    # 遍历当前文件夹下的所有文件
    for filename in os.listdir(target_folder):
        if os.path.isfile(os.path.join(target_folder, filename)):
            print(filename)
            if filename.rsplit('.', maxsplit=1)[0].endswith('english'):
                english_docx = filename
                pass
            elif filename.rsplit('.', maxsplit=1)[0].endswith('chinese'):
                chinese_docx = filename
                pass
            pass
        pass

    # destination_file = english_docx.rsplit('.', maxsplit=1)[0] + 'composition.docx'
    # 打开Word文档
    if not all([english_docx, chinese_docx]):
        print('file not found')
        return
    doc1 = Document(english_docx)
    doc2 = Document(chinese_docx)
    # 读取文档的全部内容
    word_content1 = ''
    for para in doc1.paragraphs:
        word_content1 += para.text
        pass
    word_content2 = ''
    for para in doc2.paragraphs:
        word_content2 += para.text
        pass
    reg = r'm{1,}a{1,}r{1,}k{1,}\d{8}'
    pat = re.compile(reg, re.M)
    split_doc1 = re.split(pat, word_content1)
    split_doc2 = re.split(pat, word_content2)
    #
    composition_txt = english_docx.rsplit('.', maxsplit=1)[0] + 'composition.txt'
    with open(composition_txt, mode='w', encoding='utf-8') as file:
        for i in range(len(split_doc1)):
            print('<compare>' + '-' * 60)
            print(split_doc1[i])
            print(split_doc2[i])
            file.write(split_doc1[i])
            file.write('\n')
            file.write(split_doc2[i])
            file.write('\n')
            pass
        pass
    txt2word(composition_txt)

    pass


if __name__ == '__main__':
    task1()
    pass
