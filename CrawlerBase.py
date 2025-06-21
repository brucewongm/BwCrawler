import datetime
import logging
import os
import random
import re
from logging.handlers import RotatingFileHandler
from pprint import pprint

import requests
import time
import urllib.parse

from bs4 import BeautifulSoup
from docx import Document
from fake_useragent import UserAgent
from datetime import datetime
from collections import namedtuple
from openai import OpenAI

from CrawlerBaseExclusion import exclusion
from MyAPIKey import MY_API_KEY

DebugSwitch = 1


def sort_txt_with_date(folder_path):
    """
    获取指定目录下所有文件名包含年月日的TXT文件
    并按创建时间降序排序(最新创建的在前)
    :param folder_path: 目标文件夹路径
    :return: 排序后的(文件名, 创建时间)元组列表
    """
    if not os.path.exists(folder_path):
        return None
        pass
    valid_files = []
    reg_date = r'\d{4}-\d{2}-\d{2}'
    for filename in os.listdir(folder_path):
        if filename.lower().endswith('.txt'):
            # 检查文件名是否包含年月日格式(如2025-05-17)
            if re.findall(reg_date, filename, re.I):
                filepath = os.path.join(folder_path, filename)
                if os.path.isfile(filepath):
                    ctime = os.path.getctime(filepath)  # Windows获取创建时间
                    readable_time = datetime.fromtimestamp(ctime).strftime('%Y-%m-%d %H:%M:%S')
                    valid_files.append((filename, readable_time, ctime))
                    pass
                pass
            pass
        pass

    # 按创建时间戳降序排序
    if valid_files:
        valid_files.sort(key=lambda x: x[2], reverse=True)
        # return [(f[0], f[1]) for f in valid_files]
        return [f[0] for f in valid_files]
    else:
        return None
    pass


# def sort_txt_by_creation_time(folder_path):
#     """
#     获取指定目录下所有TXT文件并按创建时间降序排序
#     :param folder_path: 目标文件夹路径
#     :return: 排序后的文件名列表(最新创建的在前)
#     """
#     if not os.path.exists(folder_path):
#         return None
#         pass
#     # 获取TXT文件列表及创建时间
#     files = []
#     for filename in os.listdir(folder_path):
#         if filename.lower().endswith('.txt'):
#             filepath = os.path.join(folder_path, filename)
#             if os.path.isfile(filepath):
#                 ctime = os.path.getctime(filepath)  # Windows专用获取创建时间
#                 files.append((filename, ctime))
#                 pass
#             pass
#         pass
#
#     # 按创建时间降序排序
#     files.sort(key=lambda x: x[1], reverse=True)
#
#     return [f[0] for f in files]

def get_deepseek_translation(user_content=None, system_content=None, temperature=1.3):
    print('Start getting deepseek response.')
    if not system_content:
        system_content = "You are a helpful assistant and capable translator"
        pass
    client = OpenAI(api_key=MY_API_KEY, base_url="https://api.deepseek.com")
    response = client.chat.completions.create(
        model="deepseek-chat",
        messages=[
            {"role": "system", "content": system_content},
            {"role": "user", "content": user_content},
        ],
        temperature=temperature,
        stream=False)
    target = response.choices[0].message.content
    return target


def get_current_datetime_tuple():
    """
    获取当前日期时间并转换为 namedtuple 格式

    返回:
        CDateTime: 包含当前年、月、日、时、分、秒的 namedtuple
    """
    current_datetime = datetime.now()

    # 定义一个 namedtuple 来存储日期和时间
    CDateTime = namedtuple('CDateTime', ['year', 'month', 'day', 'hour', 'minute', 'second'])

    # 将当前日期和时间转换为 namedtuple
    dt_tuple = CDateTime(
        year=current_datetime.year,
        month=current_datetime.month,
        day=current_datetime.day,
        hour=current_datetime.hour,
        minute=current_datetime.minute,
        second=current_datetime.second
    )

    return dt_tuple


def get_rbc_page_datetime_tuple(time_div):
    # 定义一个 namedtuple
    UrlDateTime = namedtuple('DateTime', ['year', 'month', 'day', 'hour', 'minute'])

    if time_div:
        # 提取时间文本
        date_time_str = time_div.get_text(strip=True)
        # print("date_time_str:", date_time_str)

        # 解析时间和日期字符串
        try:
            # 使用正则表达式分割字符串
            target_result = re.split(r'[:,\s-]+', date_time_str)
            # 提取出月份、日期、年份、小时和分钟
            _, month_name, day_name, year_name, hour_name, minute_name = target_result

            # 将月份转换为数字
            month_dict = {
                'January': 1, 'February': 2, 'March': 3, 'April': 4,
                'May': 5, 'June': 6, 'July': 7, 'August': 8,
                'September': 9, 'October': 10, 'November': 11, 'December': 12
            }
            month_number = month_dict[month_name]

            # 将字符串转换为整数
            year = int(year_name)
            day = int(day_name)
            hour = int(hour_name)
            minute = int(minute_name)

            # 返回 namedtuple
            return UrlDateTime(year=year, month=month_number, day=day,
                               hour=hour, minute=minute)
        except ValueError:
            print("无法解析时间格式")
            return UrlDateTime(year='', month='', day='', hour='', minute='')

    else:
        print("未找到时间信息")
        return UrlDateTime(year='', month='', day='', hour='', minute='')
    pass


def set_logger(logger_obj, logger_file_name_with_no_suffix,
               logger_file_directory=None, time_in_filename=True,
               rotating_file_handler=True):
    """
    :param logger_obj:logging.getLogger(log_name)
    :param logger_file_name_with_no_suffix:
    :param logger_file_directory:
    :param time_in_filename:
    :param rotating_file_handler:
    :return:
    """
    """ 设置循环日志，日志文件大小超过10M时，将生成备份日志，最多保留30个日志文件"""
    if not logger_obj:
        log_name = 'temp' + moment()
        logger_obj = logging.getLogger(log_name)
        pass
    dprint('entered logger_initialization')
    # log文件名称
    _file_name = logger_file_name_with_no_suffix + '.log'
    if time_in_filename:
        # 如果log文件名需要加时间
        _file_name = logger_file_name_with_no_suffix + '-' + moment() + '.log'
    # 文件所在目录
    _log_file_directory = logger_file_directory
    if not logger_file_directory:
        _log_file_directory = os.path.join(os.getcwd(), 'log')
    # create log directory and file
    if not os.path.exists(_log_file_directory):
        os.mkdir(_log_file_directory)
    # log文件绝对路径
    _log_file_abspath = os.path.join(_log_file_directory, _file_name)
    # set getLogger object
    '''设置getLogger的默认级别，添加的handler如果没有单独设置级别
        （如下面的FileHandler，单独进行了设置），将默认使用该级别'''
    logger_obj.setLevel(logging.DEBUG)
    logger_obj.propagate = False
    # create Formatter object and bind to FileHandler object
    # formatter = logging.Formatter(
    #     ("%(asctime)s - %(filename)s[line:%(lineno)d] - %(levelname)s: %(message)s"))
    formatter = logging.Formatter("%(asctime)s - : %(message)s")
    #
    '''bind file handler'''
    # create FileHandler object and bind to getLogger object
    if rotating_file_handler:
        file_handler = RotatingFileHandler(_log_file_abspath,
                                           mode='a', maxBytes=10_000_000,
                                           backupCount=30)
    else:
        file_handler = logging.FileHandler(_log_file_abspath, mode='a')
    pass
    # level:debug,info,warning,error,critical
    file_handler.setLevel(logging.DEBUG)
    # setFormatter
    file_handler.setFormatter(formatter)
    logger_obj.addHandler(file_handler)
    pass


def download_image1(url, filename):
    try:
        # if True:
        response = requests.get(url, stream=True)
        response.raise_for_status()  # 检查请求是否成功
        with open(filename, 'wb') as f:
            for chunk in response.iter_content(1024):
                f.write(chunk)
        print(f"Downloaded \n{filename}\n successfully!")
    except requests.RequestException as e:
        # else:
        print(f"Error downloading the image {url}: {e}")


def download_image(url, filename):
    try:
        # if True:
        response = requests.get(url, stream=True)
        response.raise_for_status()  # 检查请求是否成功
        with open(filename, 'wb') as f:
            for chunk in response.iter_content(1024):
                f.write(chunk)
        print(f"Downloaded \n{filename}\n successfully!")
    except requests.RequestException as e:
        # else:
        print(f"Error occurred when downloading the image {url}: {e}")
    pass


def count_down_seconds(second: int):
    for i in range(second + 1):
        left = second - i
        print('\rCounting down hold-off time {}s'.format(str(left)), end='')
        time.sleep(1)
        pass
    print('\n')
    pass


def moment():
    return time.strftime("%Y-%m-%d_%H-%M-%S", time.localtime())


def get_weekday_month_date():
    result = []
    weekdays = ['Mon', 'Tue', 'Wed', 'Thr', 'Fri', 'Sat', 'Sun']
    # 获取今天的日期
    today = datetime.date.today()
    # 获取今天是星期几（0=星期一, 1=星期二, ..., 6=星期日）
    weekday_number = today.weekday()
    result.append(weekdays[weekday_number])
    # 格式化日期为“日 英文全名月 年”的格式
    formatted_date_full_month = today.strftime("%d %B %Y")
    date_list = formatted_date_full_month.split(' ')
    result.extend(date_list[:2])
    dprint('get weekday month date:', result)
    # 打印结果
    return result


def get_rbc_page_date(soup):
    rbc_date = ''
    span_elements = soup.find_all('span', class_=True)
    for element in span_elements:
        dprint('value of element:', element)
        class_name = element.get('class')
        dprint('value-of-class_name:', class_name)
        if 'current-date' in class_name[0]:
            rbc_date = element.get_text()
            break
        pass
    dprint('value of rbc_date:', rbc_date)
    return rbc_date


def date_match_today(time_div):
    date_match = False
    current_dt = get_current_datetime_tuple()
    page_dt = get_rbc_page_datetime_tuple(time_div)
    if all([current_dt.year == page_dt.year,
            current_dt.month == page_dt.month,
            current_dt.day == page_dt.day]):
        date_match = True
    dprint('value of date compare:', date_match)
    return date_match


def verify_http_format(url):
    result = False
    reg = r'^http[s]?://\S+'
    found_re = re.findall(reg, url, re.I)
    if found_re:
        result = True
    return result


def dprint(*args, **kw_args):
    if DebugSwitch:
        print(*args, **kw_args)
    pass


def txt2word(abs_full_txt_file_name, abs_full_word_file_name=None):
    # 读取TXT文件内容
    if not abs_full_word_file_name:
        abs_full_word_file_name = abs_full_txt_file_name.rsplit('.', maxsplit=1)[0] + '.docx'
    with open(abs_full_txt_file_name, mode='r', encoding='utf-8') as file:
        text = file.read()
        pass
    # 创建一个Document对象
    doc = Document()
    # 添加文本内容到Word文档
    doc.add_paragraph(text)
    # 保存Word文档
    doc.save(abs_full_word_file_name)
    pass


def pause(seconds: int):
    seconds = random.choice(range(1, seconds))
    for i in range(seconds, 0, -1):
        time.sleep(1)
        print('\rpause counting down {}s'.format(i), end='')
        pass
    print('\rpause counting down 0s')
    pass


def compose_english_chinese_doc_files():
    print('Composing doc files...')
    current_folder = os.getcwd()
    # parent_folder = os.path.dirname(current_folder)
    # target_folder = os.path.join(parent_folder, 'results')
    today = time.strftime("%Y-%m-%d", time.localtime())
    target_folder = os.path.join(os.getcwd(), 'results{}'.format(today))
    print('target folder:{}'.format(target_folder))
    #
    english_docx = ''
    chinese_docx = ''
    # 遍历当前文件夹下的所有文件
    for filename in os.listdir(target_folder):
        this_file_path = os.path.join(target_folder, filename)
        dprint('this full filename:', this_file_path)
        if os.path.isfile(this_file_path):
            dprint('this full filename:', this_file_path)
            prename, suffix = filename.rsplit('.', maxsplit=1)
            if prename.lower().endswith('english') and suffix in 'docx':
                english_docx = this_file_path
                print('Found English file:', english_docx)
                pass
            elif prename.lower().endswith('chinese') and suffix in 'docx':
                chinese_docx = this_file_path
                print('Found Chinese file:', chinese_docx)
                pass
            pass
        pass
    # 打开Word文档
    if not all([english_docx, chinese_docx]):
        print('file not found')
        return
    doc1 = Document(english_docx)
    doc2 = Document(chinese_docx)
    # 读取文档的全部内容
    word_content1 = ''
    for para in doc1.paragraphs:
        word_content1 += para.text
        pass
    word_content2 = ''
    for para in doc2.paragraphs:
        word_content2 += para.text
        pass
    reg = r'm{1,}a{1,}r{1,}k{1,}\d{4,8}'
    pat = re.compile(reg, re.M | re.I)
    split_doc1 = re.split(pat, word_content1)
    split_doc2 = re.split(pat, word_content2)
    print('length of split doc1:', len(split_doc1))
    print('length of split doc2:', len(split_doc2))
    #
    composition_txt = english_docx.rsplit('.', maxsplit=1)[0] + 'composition.txt'
    with open(composition_txt, mode='w', encoding='utf-8') as file:
        for i in range(len(split_doc1)):
            file.write(split_doc1[i].strip())
            file.write('\n')
            file.write(split_doc2[i].strip())
            file.write('\n')
            pass
        pass
    print('Finished!')
    txt2word(composition_txt)
    os.system("start {}".format(target_folder))
    pass


def get_cookies_from(response):
    cookies = response.cookies
    # 打印Cookies
    print('cookies:', cookies)
    # 或者将Cookies保存到一个cookiejar中，然后在后续请求中使用
    cookiejar = requests.cookies.RequestsCookieJar()
    for cookie in cookies:
        cookiejar.set(cookie.name, cookie.value)
        pass
    return cookiejar


class CrawlerBase(object):
    def __init__(self, target_url, result_txt_file_name=None):
        self.target_url = target_url
        self.referred_url = target_url
        self.response = None
        self.response_text = None
        self.finished_target_url_log_file = None
        self.finished_url_list = None
        dprint('setting target url:', self.target_url)
        if result_txt_file_name:
            self.relative_txt_file_name = result_txt_file_name
            pass
        else:
            self.relative_txt_file_name = 'crawl_results_{}_original.txt'.format(moment())
            pass
        self.relative_word_file_name = None
        #
        self.log_directory = None
        self.logger = logging.getLogger('CrawlerBase')
        #
        self.current_directory = os.getcwd()
        #
        self.text_result_directory = None
        self.graphic_result_directory = None
        #
        self.abs_txt_result_file_name = None
        self.abs_word_result_file_name = None
        #
        self.headers_formed = {}
        self.gotten_page = ''
        #
        self.mark_counter = 20241101
        self.brand_task_type = True
        self.the_latest_txt_file_name = ''
        request_answer = input('Is it a brand new task ? (Non-brand-new task will collect all available txt files.):')
        if request_answer.lower() == 'n':
            self.brand_task_type = False
            pass
        #

    def initiate_environment(self):
        # 设置文件路径
        self.log_directory = os.path.join(self.current_directory, 'log')
        set_logger(self.logger, moment(), self.log_directory)
        today = time.strftime("%Y-%m-%d", time.localtime())
        if not self.text_result_directory:
            self.text_result_directory = os.path.join(
                self.current_directory, 'crawler_results_{}'.format(today))
            pass
        if not self.graphic_result_directory:
            self.graphic_result_directory = os.path.join(
                self.current_directory, 'crawler_downloaded_images_{}'.format(today))
            pass
        if not self.brand_task_type:
            existed_file_list = sort_txt_with_date(self.text_result_directory)
            if existed_file_list:
                self.abs_txt_result_file_name = os.path.join(self.text_result_directory, existed_file_list[0])
            else:
                self.abs_txt_result_file_name = os.path.join(self.text_result_directory, self.relative_txt_file_name)
                pass
            pass
        else:
            self.abs_txt_result_file_name = os.path.join(self.text_result_directory, self.relative_txt_file_name)
            pass
        self.logger.debug('initiate text file name successfully.[{}]'.format(
            self.abs_txt_result_file_name))
        #
        self.relative_word_file_name = self.relative_txt_file_name.split('.')[0] + '.docx'
        self.abs_word_result_file_name = os.path.join(self.text_result_directory, self.relative_word_file_name)
        self.logger.debug('initiate word file name successfully.[{}]'.format(
            self.abs_word_result_file_name))
        if not os.path.exists(self.text_result_directory):
            os.makedirs(self.text_result_directory)
            pass
        self.logger.debug('initiate text result folder successfully.')
        # if not os.path.exists(self.graphic_result_directory):
        #     os.makedirs(self.graphic_result_directory)
        #     pass
        # self.logger.debug('initiate picture result folder successfully.')
        # log file
        self.finished_target_url_log_file = os.path.join(self.text_result_directory, 'finished_target_url.txt')
        if not os.path.exists(self.finished_target_url_log_file):
            # 文件不存在，创建文件
            with open(self.finished_target_url_log_file, 'w') as file:
                file.write('')
                pass
            pass
        self.logger.debug('initiate finished text file successfully.')
        # 获取文件内容
        with open(self.finished_target_url_log_file, 'r+', encoding='utf-8') as file_object:
            url_lines = file_object.readlines()
            file_object.close()
            pass
        self.finished_url_list = [_.strip() for _ in url_lines]
        print('Finished urls:')
        pprint(self.finished_url_list)
        print('the number of finished url links is {}'.format(str(len(self.finished_url_list))))
        # 设置请求头，模拟正常用户的浏览器请求
        ua = UserAgent()
        try:
            self.headers_formed = {'User-Agent': ua.random}
            pass
        except:
            self.headers_formed = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/96.0.4664.45 Safari/537.36',
                'Accept-Language': 'en-US,en;q=0.5',
                'Accept-Encoding': 'gzip, deflate, br',
                'Connection': 'keep-alive',
                'Upgrade-Insecure-Requests': '1',
                'DNT': '1',  # Do Not Track 请求头
                'Referer': f'{self.referred_url}',  # 引用页，有助于避免被识别为爬虫（有时）
            }
            pass
        pass

    def set_header(self, heads):
        self.headers_formed = heads
        pass

    def set_referred_url(self, referred_url):
        self.referred_url = referred_url
        pass

    def requests_get(self):
        self.initiate_environment()
        self.response = requests.get(self.target_url, headers=self.headers_formed)
        self.response.raise_for_status()  # 检查请求是否成功
        self.response_text = self.response.text
        print('requests get finished!')
        pass

    def generate_mark(self):
        this_mark = 'mmmaaarrrkkk' + str(self.mark_counter)
        self.mark_counter += 1
        return this_mark

    def extract_response_title_link(self):
        print('extracting title and link...')
        title_link_tuple_list = []
        # 检查请求是否成功
        if self.response.status_code == 200:
            print("response.status_code:", 200)
            # 使用BeautifulSoup解析网页内容
            soup = BeautifulSoup(self.response.text, 'html.parser')
            links = soup.find_all('a')
            for link in links:
                # 获取链接的文本和URL
                gotten_text = link.get_text().strip()
                gotten_text = re.sub(r'(\s{2,})', '', gotten_text, re.M | re.I)
                gotten_link = link.get('href')
                if not verify_http_format(gotten_link):
                    continue
                elif gotten_link in exclusion:
                    continue
                else:
                    title_link_tuple_list.append((gotten_text, gotten_link))
                    pass
                pass

            pass
        else:
            print(f"Error: {self.response.status_code}")
            pass
        return title_link_tuple_list

    def extract_response_content(self, response, result_abs_txt_file_name=None, add_mark=False):
        if not result_abs_txt_file_name:
            result_abs_txt_file_name = self.abs_txt_result_file_name
        dprint('result file path:', result_abs_txt_file_name)
        # 检查请求是否成功
        if response.status_code == 200:
            dprint('extracting content...')
            soup = BeautifulSoup(response.text, 'html.parser')
            with open(result_abs_txt_file_name, mode='a+', encoding='utf-8') as file:
                for element in soup.find_all(['h1', 'h2', 'h3', 'p', 'ul', 'li']):
                    dprint('element:', element)
                    if element.name in ['h1', 'h2', 'h3']:
                        dprint('writing title to file...')
                        element_content = element.text.strip()
                        if element_content:
                            file.write(element_content)
                            file.write('\n\n')
                            pass
                        if add_mark:
                            file.write(self.generate_mark())
                            file.write('\n\n')
                            pass
                        file.flush()
                        pass
                    elif element.name == 'p':
                        dprint('writing paragraph to file...')
                        element_content = element.text.strip()
                        if element_content:
                            file.write(element_content)
                            file.write('\n\n')
                            pass
                        if add_mark:
                            file.write(self.generate_mark())
                            file.write('\n\n')
                            pass
                        file.flush()
                        pass
                    pass
                file.write('\n')
                file.flush()
                pass
            pass
        else:
            print(f"Error: {response.status_code}")
            pass
        pass

    def extract_images_list(self, base_url):
        soup = BeautifulSoup(self.response_text, 'html.parser')
        image_url_list = []
        # 找到所有的 <img> 标签
        for img_tag in soup.find_all('img'):
            # 获取图片的 src 属性
            img_url = img_tag.get('src')
            # print('>' * 100)
            # print('this image source:', img_url)
            # 如果 src 是相对路径，则将其转换为绝对路径
            if not urllib.parse.urlparse(img_url).netloc:
                img_url = urllib.parse.urljoin(base_url, img_url)
                pass
            # 将图片 URL 添加到列表中
            image_url_list.append(img_url)
        return image_url_list

    # def download_page_pictures(self):
    #     if self.response_text:
    #         base_url = "{0.scheme}://{0.netloc}".format(urllib.parse.urlparse(self.target_url))
    #         image_list = self.extract_images_list(base_url)
    #         # 创建一个文件夹来保存图片
    #         save_folder = 'downloaded_images'
    #         os.makedirs(save_folder, exist_ok=True)
    #         # 下载每张图片
    #         for idx, img_url in enumerate(image_list):
    #             print('\n' + '>' * 100)
    #             print('downloading ', img_url)
    #             # 构造文件名，例如 image1.jpg, image2.png 等
    #             file_ext = img_url.split('.')[-1]  # 获取文件扩展名
    #             if file_ext not in ['jpg', 'jpeg']:
    #                 print('this picture is not jpg or jpeg files')
    #                 continue
    #             filename = os.path.join(save_folder, f'image_{idx + 1}.{file_ext}')
    #             download_image(img_url, filename)
    #             pass
    #         pass
    #     pass
    #
    pass


class WebpagePictureDownloader(object):
    @classmethod
    def fetch_page_content(cls, url):
        try:
            response = requests.get(url)
            response.raise_for_status()  # 检查请求是否成功
            return response.text
        except requests.RequestException as e:
            print(f"Error fetching the page: {e}")
            return None
        pass

    @classmethod
    def parse_page_for_image_list(cls, content, base_url):
        soup = BeautifulSoup(content, 'html.parser')
        images = []

        # 找到所有的 <img> 标签
        for img_tag in soup.find_all('img'):
            # 获取图片的 src 属性
            img_url = img_tag.get('src')
            print('\n' + '>' * 100)
            print('this image source:', img_url)
            # 如果 src 是相对路径，则将其转换为绝对路径
            if not urllib.parse.urlparse(img_url).netloc:
                img_url = urllib.parse.urljoin(base_url, img_url)

            # 将图片 URL 添加到列表中
            images.append(img_url)

        return images

    @classmethod
    def download_one_image(cls, url, filename):
        # try:
        if True:
            dprint('Requesting this image url...\r')
            response = requests.get(url, stream=True)
            response.raise_for_status()  # 检查请求是否成功
            with open(filename, 'wb') as f:
                for chunk in response.iter_content(1024):
                    f.write(chunk)
            print(f"Downloaded {filename}")
        # except requests.RequestException as e:
        else:
            print(f"Error downloading the image {url}: {e}")

    @classmethod
    def download_webpage_pictures(cls, url: str, save_folder='downloaded_images'):
        # url = "https://newsukraine.rbc.ua/news/ukrainian-partisans-locate-strategic-military-1731456813.html"
        page_content = cls.fetch_page_content(url)
        if page_content:
            base_url = "{0.scheme}://{0.netloc}".format(urllib.parse.urlparse(url))
            image_link_list = cls.parse_page_for_image_list(page_content, base_url)
            # 创建一个文件夹来保存图片
            # save_folder = 'downloaded_images'
            os.makedirs(save_folder, exist_ok=True)
            # 下载每张图片
            url_without_suffix = url.rsplit('.', maxsplit=1)[0]
            # temp_name_list = re.split(pattern=r'[/_\-]', string=url_without_suffix, flags=re.I)
            raw_file_name = url_without_suffix.rsplit('/')[-1]
            for idx, img_url in enumerate(image_link_list):
                print('\n' + '>' * 100)
                print('Trying downloading picture from link:', img_url)
                # 构造文件名，例如 image1.jpg, image2.png 等
                file_ext = img_url.split('.')[-1]  # 获取文件扩展名
                if file_ext not in ['jpg', 'jpeg']:
                    continue
                    pass
                filename = os.path.join(save_folder, f'{raw_file_name}_image_{idx + 1}.{file_ext}')
                print('Downloading picture with name:', filename)
                cls.download_one_image(img_url, filename)
                # time.sleep(3)
                count_down_seconds(3)
                pass
            pass
        dprint('Download pictures of page \n{}\nsuccessfully!'.format(url))
        pass

    @classmethod
    def download_webpage_pictures_of_the_size(cls, url, size='650x410', save_folder='downloaded_images'):
        print('Fetching page content before downloading pictures...')
        page_content = cls.fetch_page_content(url)
        if page_content:
            base_url = "{0.scheme}://{0.netloc}".format(urllib.parse.urlparse(url))
            image_link_list = cls.parse_page_for_image_list(page_content, base_url)
            # 创建一个文件夹来保存图片
            if not os.path.exists(save_folder):
                os.makedirs(save_folder, exist_ok=True)
                pass
            # 下载每张图片
            url_without_suffix = url.rsplit('.', maxsplit=1)[0]
            # temp_name_list = re.split(pattern=r'[/_\-]', string=url_without_suffix, flags=re.I)
            raw_file_name = url_without_suffix.rsplit('/')[-1]
            for index, img_url in enumerate(image_link_list):
                print('\n' + '>' * 100)
                found_this_size = re.findall(r'\S+_(\d+x\d+)\.jpg', img_url, re.I)
                if not found_this_size:
                    continue
                else:
                    this_size = found_this_size[0]
                    pass
                print('The size of this picture probably is:{}'.format(this_size))
                print('Target picture size is {}'.format(size))
                if not size in img_url:
                    print('This picture will not be downloaded!')
                    continue
                print('Trying downloading picture from link:', img_url)
                # 构造文件名，例如 image1.jpg, image2.png 等
                file_ext = img_url.split('.')[-1]  # 获取文件扩展名
                if file_ext not in ['jpg', 'jpeg']:
                    continue
                    pass
                filename = os.path.join(save_folder, f'{raw_file_name}_image_{index + 1}.{file_ext}')
                print('Downloading picture with name:', filename)
                cls.download_one_image(img_url, filename)
                # time.sleep(6)
                count_down_seconds(6)
                pass
        pass

    pass


class BatchCrawler(CrawlerBase):
    def __init__(self, target_url, result_txt_file_name=None, crawl_number=1):
        super().__init__(target_url, result_txt_file_name)
        self.link_text_link_url_tuple_list = None
        self.response = requests.get(self.target_url, headers=self.headers_formed)
        self.crawl_number = crawl_number

    def run(self):
        self.link_text_link_url_tuple_list = self.extract_response_title_link()
        #
        with open(self.abs_txt_result_file_name, 'w') as file:
            file.write('')
            pass
        counter = 0
        for link_text, link_url in self.link_text_link_url_tuple_list:
            counter += 1
            print('round:', counter, 'dealing:', link_url)
            # print(link_text, link_url)
            # continue
            sub_response = requests.get(link_url, headers=self.headers_formed)
            self.extract_response_content(sub_response, self.abs_txt_result_file_name)
            # time.sleep(5)
            count_down_seconds(6)
            if counter >= self.crawl_number:
                break
                pass
            pass
        txt2word(self.abs_txt_result_file_name, self.abs_word_result_file_name)
        pass

    pass


pass


def task1():
    # target_url = 'https://newsukraine.rbc.ua/'
    target_url = "https://newsukraine.rbc.ua/news/ukrainian-partisans-locate-strategic-military-1731456813.html"
    crawl_number = 15
    ins = BatchCrawler(target_url, '', crawl_number)
    ins.requests_get()
    # ins.run()
    ins.download_page_pictures()
    pass


def task2():
    # count_down_seconds(6)
    ins = CrawlerBase('http')
    ins.initiate_environment()
    pass


def task3():
    directory = r"G:\MineLess\Python\projects\bwcrawler\results2025-05-17"
    pprint(sort_txt_by_creation_time(directory))
    pprint(sort_txt_with_date(directory))
    pass


if __name__ == '__main__':
    # task1()
    # task2()
    task3()
    pass
