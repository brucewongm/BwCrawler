import os
import re
from pprint import pprint

from PyQt5.QtCore import fixed

from CrawlerBase import moment
from docx import Document
from openai import OpenAI

from MyAPIKey import MY_API_KEY



def get_deepseek_response_message_content(content):
    response = client.chat.completions.create(
        model="deepseek-chat",
        messages=[
            {"role": "system", "content": "You are a helpful assistant and capable translator"},
            {"role": "user", "content": content},
        ],
        stream=False)
    target = response.choices[0].message.content
    return target


def get_result_directory():
    """
    在当前目录下查找第一个名称匹配 'resultsYYYY-MM-DD' 格式的文件夹。

    Returns:
        str | None: 第一个匹配的文件夹名（如果找到），否则返回 None。
    """
    pattern = r'results\d{4}-\d{1,2}-\d{1,2}'  # 硬编码正则表达式
    current_dir = os.getcwd()
    for entry in os.listdir(current_dir):
        if os.path.isdir(os.path.join(current_dir, entry)) and re.fullmatch(pattern, entry):
            return os.path.join(current_dir, entry)  # 找到即返回
    return None  # 未找到


def get_result_txt_file_full_name():
    result_directory = get_result_directory()
    assert result_directory
    if result_directory:
        pattern = r'crawl_result\d{4}-\d{1,2}-\d{1,2}_\d{1,2}-\d{1,2}-\d{1,2}_english\.txt'
        for root, _, files in os.walk(result_directory):
            for file in files:
                if file.endswith('.txt') and re.fullmatch(pattern, file):
                    return os.path.join(root, file)  # 找到第一个匹配项后立即返回
                pass
            pass
        return None  # 没有找到匹配项
        pass
    return


def get_image_directory():
    """递归搜索当前目录及其子目录"""
    pattern = r'downloaded_images\d{4}-\d{2}-\d{2}'
    for root, dirs, _ in os.walk(os.getcwd()):
        for dir_name in dirs:
            if re.fullmatch(pattern, dir_name):
                return os.path.join(root, dir_name)
    return None


def get_image_filenames():
    target_path = get_image_directory()
    file_list = []
    # 检查路径是否存在
    if not os.path.exists(target_path):
        raise FileNotFoundError(f"路径不存在: {target_path}")

    # 检查是否是目录
    if not os.path.isdir(target_path):
        raise NotADirectoryError(f"提供的路径不是目录: {target_path}")

    # 遍历目录获取所有文件
    for entry in os.listdir(target_path):
        full_path = os.path.join(target_path, entry)
        if os.path.isfile(full_path):  # 只添加文件，不包含目录
            file_list.append(entry)

    return file_list


def get_image_file_name_keyword(file_name):
    """
    :param file_name: 这个文件明会带有下划线和横线已经数字后缀，需要将这些字符进行清除
    :return:
    """
    file_name_key_word_list = re.split(r'-\d+_', file_name)
    if file_name_key_word_list:
        target = file_name_key_word_list[0]
        target = ' '.join(target.split('-'))
        return target
    return None


# print(response.choices[0].message.content)
def check1():
    target = get_deepseek_response_message_content('hello')
    print(target)
    pass


def check2():
    count = 3
    file = open('craw_result2025-04-18_20-33-20_english.txt', mode="r+", encoding="utf-8")
    content = file.read().strip()
    paragraph = re.split(r'title:\s+', content)
    for i in range(count):
        print(str(i) * 80)
        print(paragraph[i])
        question = "translate the following content into Chinese:" + paragraph[i]
        res = get_deepseek_response_message_content(question)
        print(res)
        pass
    pass


def process_result():
    target_folder = get_result_directory()
    target_txt = get_result_txt_file_full_name()
    assert target_txt
    target_image_folder = get_image_directory()
    assert target_image_folder
    output_path = os.path.join(target_folder, 'translated' + moment() + '.docx')
    print('target word path:', output_path)
    # x = input(":")
    target_image_list = get_image_filenames()
    #
    content = open(target_txt, 'r+', encoding='utf-8').read()
    content = re.sub(r'\s+title:\s+title:\s+', '', content)
    content = content.strip()
    # print(content)
    content_list = content.split('title:')
    # print(content_list)
    print('content length:', len(content_list))
    doc = Document()
    paragraph_count = 0
    for paragraph in content_list:
        paragraph_count += 1
        print('current paragraph index:', paragraph_count)
        # add text
        prompt = ('translate the following content into Chinese;'
                  'for those names in parentheses,delete the names and the parentheses;'
                  'in order to make the translation in a good-looking format,delete some certain specific symbols;'
                  'the content is as follows:')
        question = prompt + paragraph
        print('question:\n', question)
        translated_paragraph = get_deepseek_response_message_content(question)
        print('deepseek answer:\n', translated_paragraph)
        doc.add_paragraph(translated_paragraph)
        # add picture
        for image_file_name in target_image_list:
            fixed_file_name = get_image_file_name_keyword(image_file_name)
            if not fixed_file_name:
                continue
            if fixed_file_name in paragraph:
                # 将图片英文名加入段落
                doc.add_paragraph(image_file_name)
                print('found matched image:', image_file_name)
                full_image_path = os.path.join(target_image_folder, image_file_name)
                last_paragraph = doc.add_paragraph()
                run = last_paragraph.add_run()
                run.add_picture(full_image_path)
                break
                pass
            pass

        pass
    # 保存文档
    doc.save(output_path)
    return


if __name__ == '__main__':
    process_result()
    pass
