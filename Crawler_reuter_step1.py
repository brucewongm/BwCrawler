import os
import time
from urllib.parse import urljoin
from pprint import pprint

import requests
from bs4 import BeautifulSoup
from translate import Translator
import re
from docx import Document

OPEN = 1
CLOSED = 0


def moment():
    return time.strftime("%Y-%m-%d_%H-%M-%S", time.localtime())


def english_to_chinese(text):
    translator = Translator(from_lang="en", to_lang="zh")
    reg = r'\s{2,}'
    pat = re.compile(reg, re.I | re.M)
    result = ''
    split_parts = re.split(pat, text)
    print(split_parts)
    for line in split_parts:
        line = line.strip()
        print('line:', line)
        print('mark start')
        if not line:
            print('mark empty')
            continue
            pass
        if len(line) >= 500:
            sentence_list = []
            second_units = line.split(',')
            for sentence in second_units:
                second_translation = translator.translate(sentence)
                sentence_list.append(second_translation)
                pass
            result += ','.join(sentence_list)
            pass
        else:
            translation = translator.translate(line)
            result += translation + '\n'
            pass
    return result


def txt2word(abs_full_txt_file_name, abs_full_word_file_name=None):
    # 读取TXT文件内容
    if not abs_full_word_file_name:
        abs_full_word_file_name = abs_full_txt_file_name.rsplit('.', maxsplit=1)[0] + '.docx'
    with open(abs_full_txt_file_name, mode='r', encoding='utf-8') as file:
        text = file.read()
        pass
    # 创建一个Document对象
    doc = Document()
    # 添加文本内容到Word文档
    doc.add_paragraph(text)
    # 保存Word文档
    doc.save(abs_full_word_file_name)
    pass


def pause(seconds: int):
    for i in range(seconds, 0, -1):
        time.sleep(1)
        print('\rpause counting down {}s'.format(i), end='')
        pass
    print('\rpause counting down 0s')
    pass


def compose_result_folder_doc_files():
    current_folder = os.getcwd()
    target_folder = os.path.join(current_folder, 'results')
    english_docx = ''
    chinese_docx = ''
    # 遍历当前文件夹下的所有文件
    for filename in os.listdir(target_folder):
        if os.path.isfile(os.path.join(target_folder, filename)):
            print(filename)
            if filename.rsplit('.', maxsplit=1)[0].endswith('english'):
                english_docx = filename
                pass
            elif filename.rsplit('.', maxsplit=1)[0].endswith('chinese'):
                chinese_docx = filename
                pass
            pass
        pass

    # destination_file = english_docx.rsplit('.', maxsplit=1)[0] + 'composition.docx'
    # 打开Word文档
    if not all([english_docx, chinese_docx]):
        print('file not found')
        return
    doc1 = Document(english_docx)
    doc2 = Document(chinese_docx)
    # 读取文档的全部内容
    word_content1 = ''
    for para in doc1.paragraphs:
        word_content1 += para.text
        pass
    word_content2 = ''
    for para in doc2.paragraphs:
        word_content2 += para.text
        pass
    reg = r'm{1,}a{1,}r{1,}k{1,}\d{8}'
    pat = re.compile(reg, re.M)
    split_doc1 = re.split(pat, word_content1)
    split_doc2 = re.split(pat, word_content2)
    #
    composition_txt = english_docx.rsplit('.', maxsplit=1)[0] + 'composition.txt'
    with open(composition_txt, mode='w', encoding='utf-8') as file:
        for i in range(len(split_doc1)):
            print('<compare>' + '-' * 60)
            print(split_doc1[i])
            print(split_doc2[i])
            file.write(split_doc1[i])
            file.write('\n')
            file.write(split_doc2[i])
            file.write('\n')
            pass
        pass
    txt2word(composition_txt)

    pass

class CrawlerOld(object):
    def __init__(self, target_url, headers, crawl_number=1, result_txt_file_name=None):
        self.headers = headers
        self.crawl_number = crawl_number
        self.target_url = target_url
        #
        if result_txt_file_name:
            self.result_txt_file_name = result_txt_file_name
            pass
        else:
            self.result_txt_file_name = 'crawl_results_{}.txt'.format(moment())
        self.result_word_file_name = self.result_txt_file_name.rsplit('.', maxsplit=1)[0] + '.docx'
        self.current_directory = os.getcwd()
        self.result_directory = os.path.join(self.current_directory, 'results')
        self.result_abs_txt_file_name = os.path.join(self.result_directory, self.result_txt_file_name)
        # print(self.result_abs_txt_file_name)
        self.result_abs_word_file_name = os.path.join(self.result_directory, self.result_word_file_name)
        # print(self.result_abs_word_file_name)
        #
        self.front_door_state = OPEN
        self.back_door_state = OPEN

    pass

    def initiate_environment(self):
        if not os.path.exists(self.result_directory):
            os.makedirs(self.result_directory)
            pass
        pass

    def crawl_main_page_urls(self):
        urls = []
        response = requests.get(self.target_url)
        # 检查请求是否成功
        if response.status_code == 200:
            # 使用BeautifulSoup解析网页内容
            soup = BeautifulSoup(response.text, 'html.parser')
            links = soup.find_all('a')
            for link in links:
                # 获取链接的文本和URL
                # gotten_link = link.get('href')
                # print(link.get_text(), link.get('href'))
                urls.append((link.get_text(), link.get('href')))
                pass

            pass
        else:
            print(f"Error: {response.status_code}")
            pass
        pass
        return urls

    def record_available(self):
        if self.front_door_state == CLOSED and self.back_door_state == OPEN:
            result = True
            pass
        else:
            result = False
            pass
        return result

    def crawl_url_content(self, link_text, link_url, filename):
        self.front_door_state = OPEN
        self.back_door_state = OPEN
        print('dealing', link_url)
        response = requests.get(link_url, headers=self.headers)
        # 检查请求是否成功
        target_list = []
        if response.status_code == 200:
            # 使用BeautifulSoup解析网页内容
            soup = BeautifulSoup(response.text, 'html.parser')
            body_text = soup.body.get_text()
            print(body_text)
            #
            with open(filename, mode='a+', encoding='utf-8') as file1:
                # write headline
                file1.write(link_text + '\n')
                #
                # file1.write(target_content)
                file1.write(body_text)
                file1.write('\n\n')
                file1.flush()
                pass
            pass
        else:
            print(f"Error: {response.status_code}")
            pass
        pass

    def loop_crawl(self):
        # url_list = open('urls.txt').read().splitlines()
        # pprint(url_list)
        link_text_link_url_tuple_list = self.crawl_main_page_urls()
        #
        with open(self.result_abs_txt_file_name, 'w') as file:
            file.write('')
            pass
        counter = 0
        for link_text, link_url in link_text_link_url_tuple_list:
            counter += 1
            print('round:', counter)
            self.crawl_url_content(link_text, link_url, self.result_abs_txt_file_name)
            if counter >= self.crawl_number:
                break
                pass
            else:
                pause(8)
                pass
            pass
        pass

    def txt2word(self):
        # # 读取TXT文件内容
        # with open(self.result_abs_txt_file_name, mode='r', encoding='utf-8') as file:
        #     text = file.read()
        #     pass
        # # 创建一个Document对象
        # doc = Document()
        # # 添加文本内容到Word文档
        # doc.add_paragraph(text)
        # # 保存Word文档
        # doc.save(self.result_abs_word_file_name)
        txt2word(self.result_abs_txt_file_name, self.result_word_file_name)
        pass

    def run(self):
        self.initiate_environment()
        self.loop_crawl()
        self.txt2word()
        pass

    pass


pass


class CrawlReuters:
    def __init__(self):
        self.mark_counter = 20241101
        pass

    def generate_mark(self):
        this_mark = 'mmmaaarrrkkk' + str(self.mark_counter)
        self.mark_counter += 1
        return this_mark

    def run(self):
        # 主页的URL
        main_page_url = 'https://www.reuters.com'
        relative_url = '/world/ukraine-russia-war/'
        full_url = urljoin(main_page_url, relative_url)
        #
        result_txt_file_name = "reuters_{}english.txt".format(moment())
        current_directory = os.getcwd()
        result_directory = os.path.join(current_directory, 'results')
        result_abs_txt_file_name = os.path.join(result_directory, result_txt_file_name)
        if not os.path.exists(result_directory):
            os.makedirs(result_directory)
            pass
        pass

        # 设置请求头，模拟正常用户的浏览器请求
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/96.0.4664.45 Safari/537.36',
            'Accept-Language': 'en-US,en;q=0.5',
            'Accept-Encoding': 'gzip, deflate, br',
            # 'Connection': 'keep-alive',
            'Connection': 'close',
            'Upgrade-Insecure-Requests': '1',
            'DNT': '1',  # Do Not Track 请求头
            'Referer': main_page_url,  # 引用页，有助于避免被识别为爬虫（有时）
        }
        # print('make-first-get,' * 10)
        pause(3)
        print('start requests getting......')
        # 获取cookies
        response_main_page = requests.get(main_page_url,headers=headers)
        # 获取Cookies
        cookies = response_main_page.cookies
        # 打印Cookies
        print('cookies:', cookies)
        # 或者将Cookies保存到一个cookiejar中，然后在后续请求中使用
        cookiejar = requests.cookies.RequestsCookieJar()
        for cookie in cookies:
            cookiejar.set(cookie.name, cookie.value)
            pass
        pass
        return
        # print('make-second-get,' * 10)
        pause(10)
        # ------------------------------------------------------------------------------------
        response = requests.get(full_url, headers=headers, cookies=cookiejar)
        print("response.status_code:")
        print(response.status_code)
        soup = BeautifulSoup(response.text, 'html.parser')
        #
        with open(result_abs_txt_file_name, mode='w', encoding='utf-8') as file:
            for element in soup.find_all(['h1', 'h2', 'h3', 'p', 'ul', 'li']):
                if element.name in ['h1', 'h2', 'h3', 'p']:
                    file.write(element.text.strip())
                    file.write('\n')
                    file.write(self.generate_mark())
                    file.write('\n')
                    file.flush()
                    pass
                elif element.name == 'ul':
                    list_items = element.find_all('li')
                    for li_element in list_items:
                        file.write('!!!:' + li_element.text.strip())
                        file.write('\n')
                        file.write(self.generate_mark())
                        file.write('\n')
                        file.flush()
                        pass
                    pass
                pass
            pass
        #
        txt2word(result_abs_txt_file_name, result_abs_txt_file_name.split('.')[0] + '.docx')
        pause(30)
        print('crawling finished!')
        pass

    pass


pass


def task1():
    ins = CrawlReuters()
    ins.run()
    pass



if __name__ == '__main__':
    task1()
    pass
