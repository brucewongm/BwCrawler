# import os
import os.path
from typing import List
from urllib.parse import urljoin
# import requests
# from bs4 import BeautifulSoup
import chardet
from docx.shared import Inches
# from tqdm import tqdm
from concurrent.futures import ThreadPoolExecutor
from CrawlerBase import *


# class ISWScraper:
#     def __init__(self, base_url="https://understandingwar.org",
#                  image_dir="isw_downloaded_images_" + moment()):
#         self.base_url = base_url
#         self.image_dir = image_dir
#         self.headers = {
#             'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
#         }
#         self._setup_directories()
#
#     def _setup_directories(self):
#         """创建必要的目录"""
#         os.makedirs(self.image_dir, exist_ok=True)
#
#     def download_image(self, url):
#         """下载单张图片"""
#         try:
#             response = requests.get(url, stream=True)
#             response.raise_for_status()
#             filename = os.path.basename(url.split('?')[0])
#             save_path = os.path.join(self.image_dir, filename)
#
#             with open(save_path, 'wb') as f:
#                 for chunk in response.iter_content(chunk_size=8192):
#                     f.write(chunk)
#             print(f"下载成功: {filename}")
#             return True
#         except Exception as e:
#             print(f"下载失败 {url}: {e}")
#             return False
#
#     def download_pdf(self, pdf_link, destination_directory=None, base_file_name=None):
#         """下载PDF文件"""
#         try:
#             response = requests.get(pdf_link, stream=True)
#             response.raise_for_status()
#             if not destination_directory:
#                 destination_directory = os.getcwd()
#                 pass
#             if not base_file_name:
#                 base_file_name = os.path.basename(pdf_link)
#                 pass
#             filename = os.path.join(destination_directory, base_file_name)
#             with open(filename, 'wb') as f:
#                 for chunk in response.iter_content(chunk_size=8192):
#                     f.write(chunk)
#                     pass
#                 pass
#             print(f"PDF文件已成功下载到: {os.path.abspath(filename)}")
#             return True
#         except Exception as e:
#             print(f"下载PDF时出错: {e}")
#             return False
#
#     def _extract_metadata(self, soup):
#         """提取文章元数据"""
#         title = soup.find('h1', {'id': 'page-title'}).get_text(strip=True)
#         submitted = soup.find('span', class_='submitted').get_text(strip=True)
#
#         print("Title:", title)
#         print("Publication Info:", submitted)
#
#         return {'title': title, 'submitted': submitted}
#
#     def _extract_main_content(self, soup):
#         """提取正文内容"""
#         content_div = soup.find('div', class_='field-name-body')
#         if not content_div:
#             print("Could not find main content")
#             return ''
#
#         paragraphs = [p.get_text(strip=True) for p in content_div.find_all('p')]
#         main_content = '\n'.join(paragraphs)
#
#         print("\nMain Content:")
#         print(main_content)
#
#         return main_content
#
#     def _extract_pdf_link(self, soup):
#         """提取PDF链接"""
#         pdf_div = soup.find('div', class_='field-name-field-pdf-report')
#         if not pdf_div:
#             return ''
#         pdf_link = pdf_div.find('a')['href']
#         if not pdf_link.startswith('http'):
#             pdf_link = urljoin(self.base_url, pdf_link)
#             pass
#
#         print("\nThis PDF Link:", pdf_link)
#         return pdf_link
#
#     def _extract_all_content(self, soup: BeautifulSoup):
#         title = soup.find('h1', {'id': 'page-title'}).get_text(strip=True)
#         submitted = soup.find('span', class_='submitted').get_text(strip=True)
#         pdf_link = self._extract_pdf_link(soup)
#         contents = []
#         tags = ['h1', 'h2', 'h3', 'p', 'ul', 'li']
#         elements = soup.find_all(tags)
#         for element in elements:
#             contents.append(element.get_text(strip=True))
#             pass
#
#         # elements = soup.find_all(lambda tag:
#         #                          (tag.name in tags) or
#         #                          (tag.name == 'h1' and tag.get('id') == 'page-title') or
#         #                          (tag.name == 'span' and 'submitted' in tag.get('class', [])))
#         return {'title': title, 'submitted': submitted, 'contents': '\n'.join([content for content in contents])}
#
#     def _extract_and_download_images(self, soup):
#         """提取并下载图片"""
#         content_div = soup.find('div', class_='field-name-body')
#         if not content_div:
#             print("未找到正文内容")
#             return []
#
#         images = content_div.find_all('img')
#         print(f"找到 {len(images)} 张图片")
#
#         image_urls = []
#         for img in images:
#             img_url = img.get('src')
#             if not img_url.startswith(('http://', 'https://')):
#                 img_url = urljoin(self.base_url, img_url)
#
#             # 只下载特定类型的图片
#             valid_extensions = ['.jpg', '.jpeg', '.png']
#             if not any(img_url.lower().endswith(ext) for ext in valid_extensions):
#                 continue
#
#             image_urls.append(img_url)
#             pass
#
#         print('开始多线程下载图片...')
#         with ThreadPoolExecutor(max_workers=5) as executor:
#             executor.map(self.download_image, image_urls)
#
#         return image_urls
#
#     def process_image_element(self, img_element):
#         img_url = img_element.get('src')
#         if not img_url.startswith(('http://', 'https://')):
#             img_url = urljoin(self.base_url, img_url)
#
#         # 只下载特定类型的图片
#         valid_extensions = ['.jpg', '.jpeg', '.png']
#         if not any(img_url.lower().endswith(ext) for ext in valid_extensions):
#             return ''
#         return img_url
#
#     def _extract_all_text(self, soup):
#         # 获取原始元素列表（保持文档顺序）
#         elements = soup.find_all(lambda tag:
#                                  (tag.name in {'h1', 'h2', 'h3', 'p', 'ul', 'li'}) or
#                                  (tag.name == 'h1' and tag.get('id') == 'page-title') or
#                                  (tag.name == 'span' and 'submitted' in tag.get('class', [])) or
#                                  (tag.name == 'img' and tag.find_parent('div', class_='field-name-body'))
#                                  )
#
#         # 按原始顺序输出元素内容
#         # result = [elem.get_text(strip=True) for elem in elements]
#         #
#         result = []
#         for elem in elements:
#             if elem.name == 'img':
#                 result.append(self.process_image_element(elem))
#                 pass
#             else:
#                 # 提取文本内容
#                 result.append(elem.get_text(strip=True))
#                 pass
#             pass
#
#         return result
#
#     def scrape_article(self, url, download_images=True, download_pdf=True):
#         """主方法：抓取文章内容"""
#         try:
#             response = requests.get(url, headers=self.headers)
#             response.raise_for_status()
#
#             # 检测编码
#             detected_encoding = chardet.detect(response.content)['encoding']
#             print(f"Detected encoding: {detected_encoding}")
#
#             soup = BeautifulSoup(response.content, 'html.parser', from_encoding=detected_encoding)
#
#             # 提取各种信息
#             # metadata = self._extract_metadata(soup)
#             # content = self._extract_main_content(soup)
#             # content = self._extract_all_content(soup)
#             content = self._extract_all_text(soup)
#             print('result:')
#             count = 1
#             for item in content:
#                 print('\n' + '>' * 80 + str(count))
#                 pprint(item)
#                 count += 1
#                 pass
#
#             pdf_link = self._extract_pdf_link(soup)
#             if download_pdf and pdf_link:
#                 content['pdf_link'] = pdf_link
#                 self.download_pdf(pdf_link)
#
#             if download_images:
#                 self._extract_and_download_images(soup)
#
#             print('-' * 80)
#
#             return content
#
#         except requests.exceptions.RequestException as e:
#             print(f"Request failed: {e}")
#             return None
#         except Exception as e:
#             print(f"An error occurred: {e}")
#             return None


class ISWScraperNew:
    def __init__(self, base_url="https://understandingwar.org",
                 image_dir="isw_downloaded_images_" + datetime.now().strftime("%Y%m%d_%H%M%S"),
                 output_dir="output_docs"):
        self.base_url = base_url
        self.image_dir = image_dir
        self.output_dir = output_dir
        self.headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        self._setup_directories()

    def _setup_directories(self):
        """创建必要的目录"""
        os.makedirs(self.image_dir, exist_ok=True)
        os.makedirs(self.output_dir, exist_ok=True)

    def _translate_to_chinese(self, text):
        """将英文文本翻译为中文"""
        prompt = ("Task: Translate the following text sentence-by-sentence into Chinese while:"
                  "Requirements: "
                  "1. Preserve original formatting and punctuation "
                  "2. Maintain paragraph structures "
                  "3. Strictly adhere to provided terminology lists (if any) "
                  "4. Slight cultural adaptation "
                  "5. Make error corrections when semantically crucial"
                  "Restrictions: "
                  "- No content analysis/interpretation "
                  + text)
        if DebugSwitch:
            pprint('Sending prompt:\n' + prompt)
            pass

        try:
            if text.strip():  # 只翻译非空文本
                time.sleep(1)  # 防止翻译API限制
                # translation = self.translator.translate(text, src='en', dest='zh-cn')
                translation = get_deepseek_translation(prompt)
                if DebugSwitch:
                    print('AI response:', translation)
                    pass
                return translation
            return text
        except Exception as e:
            print(f"翻译失败: {e}")
            return text  # 翻译失败返回原文

    def _get_page_object(self, url):
        """获取页面内容并返回BeautifulSoup对象"""
        try:
            response = requests.get(url, headers=self.headers)
            response.raise_for_status()
            detected_encoding = chardet.detect(response.content)['encoding']
            return BeautifulSoup(response.content, 'html.parser', from_encoding=detected_encoding)
        except Exception as e:
            print(f"获取页面内容失败: {url} - {e}")
            return None

    def find_title_url_by_keywords(self, keywords):
        """在主页查找包含关键词的文章链接"""
        soup = self._get_page_object(self.base_url)
        if not soup:
            return []

        title_url_dict_list = []
        # 查找所有文章链接（根据实际网站结构调整选择器）
        for link in soup.select('a[href*="/backgrounder/"]'):
            title = link.get_text(strip=True)
            href = link.get('href')

            if not href or not title:
                continue

            # 检查标题是否包含所有关键词（不区分大小写）
            if all(re.search(re.escape(keyword), title, re.IGNORECASE) for keyword in keywords):
                full_url = urljoin(self.base_url, href)
                title_url_dict_list.append({'title': title, 'url': full_url})
                if DebugSwitch:
                    print(f"Found matched title: {title} - {full_url}")
                    pass
                pass
            pass

        return title_url_dict_list

    def download_image(self, url):
        """下载单张图片并返回本地路径"""
        try:
            response = requests.get(url, stream=True, headers=self.headers)
            response.raise_for_status()

            filename = os.path.basename(url.split('?')[0])
            if not any(filename.lower().endswith(ext) for ext in ['.jpg', '.jpeg', '.png', '.gif']):
                filename += '.jpg'

            save_path = os.path.join(self.image_dir, filename)

            counter = 1
            while os.path.exists(save_path):
                name, ext = os.path.splitext(filename)
                save_path = os.path.join(self.image_dir, f"{name}_{counter}{ext}")
                counter += 1

            with open(save_path, 'wb') as f:
                for chunk in response.iter_content(chunk_size=8192):
                    f.write(chunk)

            print(f"下载成功: {save_path}")
            return save_path
        except Exception as e:
            print(f"下载失败 {url}: {e}")
            return None

    def download_pdf(self, pdf_link):
        """下载PDF文件"""
        try:
            response = requests.get(pdf_link, stream=True, headers=self.headers)
            response.raise_for_status()

            filename = os.path.basename(pdf_link)
            save_path = os.path.join(self.output_dir, filename)

            with open(save_path, 'wb') as f:
                for chunk in response.iter_content(chunk_size=8192):
                    f.write(chunk)

            print(f"PDF文件已成功下载到: {save_path}")
            return save_path
        except Exception as e:
            print(f"下载PDF时出错: {e}")
            return None

    def _extract_pdf_link(self, soup):
        """提取PDF链接"""
        pdf_div = soup.find('div', class_='field-name-field-pdf-report')
        if not pdf_div:
            return ''
        pdf_link = pdf_div.find('a')['href']
        if not pdf_link.startswith('http'):
            pdf_link = urljoin(self.base_url, pdf_link)
            pass

        print("\nThis PDF Link:", pdf_link)
        return pdf_link

    def _extract_metadata(self, soup):
        """提取文章元数据"""
        title = soup.find('h1', {'id': 'page-title'}).get_text(strip=True)
        submitted = soup.find('span', class_='submitted').get_text(strip=True)

        print("Title:", title)
        print("Publication Info:", submitted)

        return {
            'title': title,
            'title_zh': self._translate_to_chinese(title),
            'submitted': submitted
        }

    def _extract_content_elements(self, soup):
        """提取正文内容元素（文本和图片）"""
        content_div = soup.find('div', class_='field-name-body')
        if not content_div:
            print("Could not find main content")
            return []

        elements = []

        for element in content_div.children:
            if element.name == 'p':
                text = element.get_text(strip=True)
                if text:
                    elements.append({
                        'type': 'text',
                        'content': text,
                        'content_zh': self._translate_to_chinese(text)
                    })
            elif element.name == 'img':
                img_url = element.get('src')
                if img_url:
                    if not img_url.startswith(('http://', 'https://')):
                        img_url = urljoin(self.base_url, img_url)
                    elements.append({
                        'type': 'image',
                        'url': img_url
                    })

        return elements

    def _download_content_images(self, content_elements):
        """下载内容中的图片并更新本地路径"""
        image_urls = [elem['url'] for elem in content_elements if elem['type'] == 'image']

        print(f'开始下载 {len(image_urls)} 张图片...')
        with ThreadPoolExecutor(max_workers=5) as executor:
            local_paths = list(executor.map(self.download_image, image_urls))

        path_index = 0
        for elem in content_elements:
            if elem['type'] == 'image':
                elem['local_path'] = local_paths[path_index]
                path_index += 1

        return content_elements

    def _generate_word_document(self, metadata, content_elements, pdf_path=None):
        """生成Word文档"""
        doc = Document()
        doc.add_heading(metadata['title_zh'], level=1)
        doc.add_paragraph(f"原始标题: {metadata['title']}")
        doc.add_paragraph(f"发布日期: {metadata['submitted']}")
        doc.add_paragraph("\n")

        for elem in content_elements:
            if elem['type'] == 'text':
                doc.add_paragraph(elem['content_zh'])
                doc.add_paragraph(f"(原文: {elem['content']})", style='Intense Quote')
                doc.add_paragraph("\n")
            elif elem['type'] == 'image' and elem.get('local_path'):
                try:
                    doc.add_picture(elem['local_path'], width=Inches(5.0))
                    doc.add_paragraph(f"图片: {os.path.basename(elem['local_path'])}")
                    doc.add_paragraph("\n")
                except Exception as e:
                    print(f"无法插入图片 {elem['local_path']}: {e}")

        if pdf_path:
            doc.add_paragraph(f"附件PDF已下载: {os.path.basename(pdf_path)}")

        doc_filename = f"{metadata['title'][:50]}.docx"
        doc_path = os.path.join(self.output_dir, doc_filename)
        doc.save(doc_path)

        print(f"Word file was generated successfully: {doc_path}")
        return doc_path

    def process_image_element(self, img_element):
        img_url = img_element.get('src')
        if not img_url.startswith(('http://', 'https://')):
            img_url = urljoin(self.base_url, img_url)

        # 只下载特定类型的图片
        valid_extensions = ['.jpg', '.jpeg', '.png']
        if not any(img_url.lower().endswith(ext) for ext in valid_extensions):
            return ''
        return img_url

    def _extract_all_text(self, soup) -> List:
        # 获取原始元素列表（保持文档顺序）
        elements = soup.find_all(
            lambda tag:
            (tag.name in {'h1', 'h2', 'h3', 'p', 'ul', 'li'}) or
            (tag.name == 'h1' and tag.get('id') == 'page-title') or
            (tag.name == 'span' and 'submitted' in tag.get('class', [])) or
            (tag.name == 'img' and tag.find_parent('div', class_='field-name-body')) or
            (tag.name == 'div' and 'field-name-field-pdf-report' in tag.get('class', []))
        )
        # 按原始顺序输出元素内容
        # result = [elem.get_text(strip=True) for elem in elements]
        #
        result = []
        for elem in elements:
            if elem.name == 'img':
                result.append(('type_image_url', self.process_image_element(elem)))
                pass
            elif elem.name == 'div' and 'field-name-field-pdf-report' in elem.get('class', []):
                link = elem.find('a', href=True)
                if link:
                    pdf_url = link['href']
                    # URL标准化处理
                    if not pdf_url.startswith(('http://', 'https://')):
                        pdf_url = urljoin(self.base_url, pdf_url)
                    result.append(('type_pdf_url', pdf_url))
                    pass
                pass
            else:
                # 提取文本内容
                result.append(('type_text', elem.get_text(strip=True)))
                pass
            pass

        return result

    def scrape_and_export_old(self, url):
        """处理单个文章URL"""
        print(f"\nStart dealing with url: {url}")
        soup = self._get_page_object(url)
        if not soup:
            return None

        metadata = self._extract_metadata(soup)
        content_elements = self._extract_content_elements(soup)
        pdf_link = self._extract_pdf_link(soup)

        pdf_path = None
        if pdf_link:
            pdf_path = self.download_pdf(pdf_link)

        content_elements = self._download_content_images(content_elements)
        doc_path = self._generate_word_document(metadata, content_elements, pdf_path)

        print('-' * 80)
        return {
            'metadata': metadata,
            'doc_path': doc_path
        }

    def scrape_and_export(self, url) -> List:
        """处理单个文章URL"""
        print(f"\nStart dealing with url: {url}")
        soup = self._get_page_object(url)
        if not soup:
            return []

        result = self._extract_all_text(soup)
        # print(result)
        return result

        # metadata = self._extract_metadata(soup)
        # content_elements = self._extract_content_elements(soup)
        # pdf_link = self._extract_pdf_link(soup)
        #
        # pdf_path = None
        # if pdf_link:
        #     pdf_path = self.download_pdf(pdf_link)
        #
        # content_elements = self._download_content_images(content_elements)
        # doc_path = self._generate_word_document(metadata, content_elements, pdf_path)
        #
        # print('-' * 80)
        # return {
        #     'metadata': metadata,
        #     'doc_path': doc_path
        # }
        pass

    def process_keyword_articles(self, keywords, number=0):
        """主方法：根据关键词查找并处理文章"""
        title_url_dicts: List = self.find_title_url_by_keywords(keywords)
        if not title_url_dicts:
            print("没有找到匹配关键词的文章")
            return []

        results = []
        count = 0
        for title_url_dict in title_url_dicts:
            print('Dealing item index:', count)
            result = self.scrape_and_export(title_url_dict['url'])
            if result:
                results.append(result)
                count += 1
                pass
            time.sleep(2)  # 礼貌性延迟
            if count == number:
                print('Target number reached ', count, '!')
                break
                pass
            pass

        return results

    pass


# def check1():
#     scraper = ISWScraper()
#
#     # 示例URL
#     url = "https://understandingwar.org/backgrounder/russian-offensive-campaign-assessment-april-28-2025/"
#
#     # 抓取文章
#     result = scraper.scrape_article(url, download_images=False, download_pdf=False)
#
#     for k, v in result.items():
#         print(k, ":", v)
#
#     # 也可以单独下载PDF
#     # pdf_link = "https://understandingwar.org/sites/default/files/2025-04-28-PDF-Russian%20Offensive%20Campaign%20Assessment.pdf"
#     # scraper.download_pdf(pdf_link)
#

def check2():
    scraper = ISWScraperNew()

    # 定义关键词列表
    title_keywords = ['russian', 'offensive', 'campaign', 'assessment']

    # 执行处理
    results = scraper.process_keyword_articles(title_keywords, 1)
    print('About to print result')
    for ele in results:
        for type_,sub_ele in ele:
            print(type_)
            pprint(sub_ele)
            pass
        pass

    # 输出结果摘要
    pass


# 使用示例
if __name__ == "__main__":
    check2()
    pass
